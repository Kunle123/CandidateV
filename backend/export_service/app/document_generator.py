import os
import asyncio
import logging
import jinja2
import tempfile
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, Optional, Tuple

# Import PDF and DOCX generation libraries
from weasyprint import HTML, CSS
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Set up logging
logger = logging.getLogger("export_service.document_generator")

# Templates directory
TEMPLATES_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "templates")
os.makedirs(TEMPLATES_DIR, exist_ok=True)

# Default template paths
DEFAULT_TEMPLATE_PATH = os.path.join(TEMPLATES_DIR, "default.html")

# Create default template if it doesn't exist
if not os.path.exists(DEFAULT_TEMPLATE_PATH):
    with open(DEFAULT_TEMPLATE_PATH, "w") as f:
        f.write("""
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{{ cv.title }}</title>
    <style>
        body {
            font-family: Arial, sans-serif;
            margin: 0;
            padding: 20px;
            color: #333;
        }
        .header {
            text-align: center;
            margin-bottom: 20px;
        }
        h1 {
            color: #2c3e50;
            margin-bottom: 5px;
        }
        .personal-info {
            text-align: center;
            margin-bottom: 20px;
        }
        .section {
            margin-bottom: 20px;
        }
        .section-title {
            border-bottom: 2px solid #3498db;
            padding-bottom: 5px;
            margin-bottom: 10px;
            color: #2c3e50;
        }
        .experience-item, .education-item {
            margin-bottom: 15px;
        }
        .item-header {
            display: flex;
            justify-content: space-between;
            margin-bottom: 5px;
        }
        .item-title {
            font-weight: bold;
        }
        .item-date {
            color: #7f8c8d;
        }
        .item-subtitle {
            font-style: italic;
            margin-bottom: 5px;
        }
        .item-description {
            margin-top: 5px;
        }
        .skills-list {
            display: flex;
            flex-wrap: wrap;
            gap: 10px;
        }
        .skill-item {
            background-color: #e6f3ff;
            padding: 5px 10px;
            border-radius: 15px;
        }
    </style>
</head>
<body>
    <div class="header">
        <h1>{{ cv.title }}</h1>
    </div>
    
    <div class="personal-info">
        <div>{{ cv.personal_info.first_name }} {{ cv.personal_info.last_name }}</div>
        <div>{{ cv.personal_info.email }} | {{ cv.personal_info.phone }}</div>
        <div>{{ cv.personal_info.location }}</div>
    </div>
    
    {% if cv.summary %}
    <div class="section">
        <h2 class="section-title">Summary</h2>
        <p>{{ cv.summary }}</p>
    </div>
    {% endif %}
    
    {% if cv.experience and cv.experience|length > 0 %}
    <div class="section">
        <h2 class="section-title">Experience</h2>
        {% for exp in cv.experience %}
        <div class="experience-item">
            <div class="item-header">
                <div class="item-title">{{ exp.title }}</div>
                <div class="item-date">
                    {{ exp.start_date }} - {% if exp.current %}Present{% else %}{{ exp.end_date }}{% endif %}
                </div>
            </div>
            <div class="item-subtitle">{{ exp.company }}{% if exp.location %}, {{ exp.location }}{% endif %}</div>
            <div class="item-description">{{ exp.description }}</div>
        </div>
        {% endfor %}
    </div>
    {% endif %}
    
    {% if cv.education and cv.education|length > 0 %}
    <div class="section">
        <h2 class="section-title">Education</h2>
        {% for edu in cv.education %}
        <div class="education-item">
            <div class="item-header">
                <div class="item-title">{{ edu.degree }} in {{ edu.field_of_study }}</div>
                <div class="item-date">
                    {{ edu.start_date }} - {{ edu.end_date }}
                </div>
            </div>
            <div class="item-subtitle">{{ edu.institution }}</div>
            {% if edu.description %}
            <div class="item-description">{{ edu.description }}</div>
            {% endif %}
        </div>
        {% endfor %}
    </div>
    {% endif %}
    
    {% if cv.skills and cv.skills|length > 0 %}
    <div class="section">
        <h2 class="section-title">Skills</h2>
        <div class="skills-list">
            {% for skill in cv.skills %}
            <div class="skill-item">{{ skill }}</div>
            {% endfor %}
        </div>
    </div>
    {% endif %}
    
    {% if cv.certifications and cv.certifications|length > 0 %}
    <div class="section">
        <h2 class="section-title">Certifications</h2>
        {% for cert in cv.certifications %}
        <div class="education-item">
            <div class="item-header">
                <div class="item-title">{{ cert.name }}</div>
                <div class="item-date">{{ cert.date }}</div>
            </div>
            <div class="item-subtitle">{{ cert.issuer }}</div>
            {% if cert.description %}
            <div class="item-description">{{ cert.description }}</div>
            {% endif %}
        </div>
        {% endfor %}
    </div>
    {% endif %}
</body>
</html>
""")

# Initialize Jinja2 environment
template_loader = jinja2.FileSystemLoader(searchpath=TEMPLATES_DIR)
template_env = jinja2.Environment(loader=template_loader)

class DocumentGenerator:
    """
    Class for generating PDF and DOCX documents from CV data
    """
    
    @staticmethod
    async def generate_pdf_weasyprint(cv_data: Dict[str, Any], output_path: str, template_options: Optional[Dict[str, Any]] = None) -> str:
        """Generate a PDF using WeasyPrint and HTML/CSS templates"""
        try:
            # Get template
            template_name = template_options.get("template_name", "default.html") if template_options else "default.html"
            template = template_env.get_template(template_name)
            
            # Render HTML
            html_content = template.render(cv=cv_data, options=template_options)
            
            # Create a temporary file for the HTML
            with tempfile.NamedTemporaryFile(suffix='.html', delete=False) as tmp_html:
                tmp_html.write(html_content.encode('utf-8'))
                tmp_html_path = tmp_html.name
            
            # Use WeasyPrint to convert HTML to PDF
            html = HTML(filename=tmp_html_path)
            
            # Add any custom styles from template options
            css_styles = []
            if template_options and "custom_css" in template_options:
                css_styles.append(CSS(string=template_options["custom_css"]))
            
            # Generate PDF
            html.write_pdf(output_path, stylesheets=css_styles)
            
            # Clean up temporary file
            os.unlink(tmp_html_path)
            
            logger.info(f"PDF generated at {output_path} using WeasyPrint")
            return output_path
            
        except Exception as e:
            logger.error(f"Error generating PDF with WeasyPrint: {str(e)}")
            raise
    
    @staticmethod
    async def generate_pdf_reportlab(cv_data: Dict[str, Any], output_path: str, template_options: Optional[Dict[str, Any]] = None) -> str:
        """Generate a PDF using ReportLab (programmatic approach)"""
        try:
            # Create PDF document
            doc = SimpleDocTemplate(
                output_path,
                pagesize=letter,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72
            )
            
            # Define styles
            styles = getSampleStyleSheet()
            title_style = styles['Title']
            heading_style = styles['Heading2']
            normal_style = styles['Normal']
            
            # Customize styles based on template options
            if template_options and "color_scheme" in template_options:
                color_scheme = template_options["color_scheme"]
                if color_scheme == "blue":
                    heading_style.textColor = colors.blue
                elif color_scheme == "green":
                    heading_style.textColor = colors.green
            
            # Build PDF content
            elements = []
            
            # Title
            elements.append(Paragraph(cv_data.get("title", "Curriculum Vitae"), title_style))
            elements.append(Spacer(1, 12))
            
            # Personal Info
            personal_info = cv_data.get("personal_info", {})
            elements.append(Paragraph(f"{personal_info.get('first_name', '')} {personal_info.get('last_name', '')}", styles['Heading1']))
            elements.append(Paragraph(f"Email: {personal_info.get('email', '')}", normal_style))
            elements.append(Paragraph(f"Phone: {personal_info.get('phone', '')}", normal_style))
            elements.append(Paragraph(f"Location: {personal_info.get('location', '')}", normal_style))
            elements.append(Spacer(1, 12))
            
            # Summary
            if "summary" in cv_data and cv_data["summary"]:
                elements.append(Paragraph("Summary", heading_style))
                elements.append(Spacer(1, 6))
                elements.append(Paragraph(cv_data["summary"], normal_style))
                elements.append(Spacer(1, 12))
            
            # Experience
            if "experience" in cv_data and cv_data["experience"]:
                elements.append(Paragraph("Experience", heading_style))
                elements.append(Spacer(1, 6))
                
                for exp in cv_data["experience"]:
                    job_title = exp.get("title", "")
                    company = exp.get("company", "")
                    start_date = exp.get("start_date", "")
                    end_date = "Present" if exp.get("current", False) else exp.get("end_date", "")
                    
                    elements.append(Paragraph(f"<b>{job_title}</b> at {company}", styles['Heading3']))
                    elements.append(Paragraph(f"{start_date} - {end_date}", normal_style))
                    elements.append(Paragraph(exp.get("description", ""), normal_style))
                    elements.append(Spacer(1, 10))
                
                elements.append(Spacer(1, 6))
            
            # Education
            if "education" in cv_data and cv_data["education"]:
                elements.append(Paragraph("Education", heading_style))
                elements.append(Spacer(1, 6))
                
                for edu in cv_data["education"]:
                    degree = edu.get("degree", "")
                    field = edu.get("field_of_study", "")
                    institution = edu.get("institution", "")
                    
                    elements.append(Paragraph(f"<b>{degree}</b> in {field}", styles['Heading3']))
                    elements.append(Paragraph(f"{institution}", normal_style))
                    elements.append(Paragraph(f"{edu.get('start_date', '')} - {edu.get('end_date', '')}", normal_style))
                    if "description" in edu and edu["description"]:
                        elements.append(Paragraph(edu["description"], normal_style))
                    elements.append(Spacer(1, 10))
                
                elements.append(Spacer(1, 6))
            
            # Skills
            if "skills" in cv_data and cv_data["skills"]:
                elements.append(Paragraph("Skills", heading_style))
                elements.append(Spacer(1, 6))
                
                # Create a skills table
                skills_data = []
                row = []
                for i, skill in enumerate(cv_data["skills"]):
                    row.append(skill)
                    if len(row) == 3 or i == len(cv_data["skills"]) - 1:
                        # Fill the row with empty cells if needed
                        while len(row) < 3:
                            row.append("")
                        skills_data.append(row)
                        row = []
                
                if skills_data:
                    skills_table = Table(skills_data, colWidths=[doc.width/3.0]*3)
                    skills_table.setStyle(TableStyle([
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                        ('BACKGROUND', (0, 0), (-1, -1), colors.lightgrey),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                        ('PADDING', (0, 0), (-1, -1), 6),
                    ]))
                    elements.append(skills_table)
                
                elements.append(Spacer(1, 12))
            
            # Build the PDF
            doc.build(elements)
            
            logger.info(f"PDF generated at {output_path} using ReportLab")
            return output_path
            
        except Exception as e:
            logger.error(f"Error generating PDF with ReportLab: {str(e)}")
            raise
    
    @staticmethod
    async def generate_docx(cv_data: Dict[str, Any], output_path: str, template_options: Optional[Dict[str, Any]] = None) -> str:
        """Generate a DOCX document using python-docx"""
        try:
            # Create a new document
            doc = Document()
            
            # Set document properties
            doc.core_properties.title = cv_data.get("title", "CV")
            doc.core_properties.author = f"{cv_data.get('personal_info', {}).get('first_name', '')} {cv_data.get('personal_info', {}).get('last_name', '')}"
            
            # Get style settings from template options
            font_name = "Calibri"
            heading_color = "0000FF"  # Blue
            
            if template_options:
                font_name = template_options.get("font", "Calibri")
                heading_color = template_options.get("heading_color", "0000FF")
            
            # Add title
            title = doc.add_heading(cv_data.get("title", "Curriculum Vitae"), 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add personal information
            personal_info = cv_data.get("personal_info", {})
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(f"{personal_info.get('first_name', '')} {personal_info.get('last_name', '')}\n").bold = True
            p.add_run(f"Email: {personal_info.get('email', '')}\n")
            p.add_run(f"Phone: {personal_info.get('phone', '')}\n")
            p.add_run(f"Location: {personal_info.get('location', '')}")
            
            doc.add_paragraph()
            
            # Add summary section
            if "summary" in cv_data and cv_data["summary"]:
                doc.add_heading("Summary", 1)
                doc.add_paragraph(cv_data["summary"])
                doc.add_paragraph()
            
            # Add experience section
            if "experience" in cv_data and cv_data["experience"]:
                doc.add_heading("Experience", 1)
                
                for exp in cv_data["experience"]:
                    job_title = exp.get("title", "")
                    company = exp.get("company", "")
                    location = exp.get("location", "")
                    start_date = exp.get("start_date", "")
                    end_date = "Present" if exp.get("current", False) else exp.get("end_date", "")
                    
                    p = doc.add_paragraph()
                    p.add_run(f"{job_title} at {company}").bold = True
                    if location:
                        p.add_run(f", {location}")
                    
                    p = doc.add_paragraph()
                    p.add_run(f"{start_date} - {end_date}")
                    
                    if "description" in exp and exp["description"]:
                        doc.add_paragraph(exp["description"])
                    
                    doc.add_paragraph()
            
            # Add education section
            if "education" in cv_data and cv_data["education"]:
                doc.add_heading("Education", 1)
                
                for edu in cv_data["education"]:
                    degree = edu.get("degree", "")
                    field = edu.get("field_of_study", "")
                    institution = edu.get("institution", "")
                    start_date = edu.get("start_date", "")
                    end_date = edu.get("end_date", "")
                    
                    p = doc.add_paragraph()
                    p.add_run(f"{degree} in {field}").bold = True
                    
                    p = doc.add_paragraph()
                    p.add_run(f"{institution}")
                    
                    p = doc.add_paragraph()
                    p.add_run(f"{start_date} - {end_date}")
                    
                    if "description" in edu and edu["description"]:
                        doc.add_paragraph(edu["description"])
                    
                    doc.add_paragraph()
            
            # Add skills section
            if "skills" in cv_data and cv_data["skills"]:
                doc.add_heading("Skills", 1)
                
                # Create a bullet list of skills
                for skill in cv_data["skills"]:
                    doc.add_paragraph(skill, style='List Bullet')
                
                doc.add_paragraph()
            
            # Add certifications section
            if "certifications" in cv_data and cv_data["certifications"]:
                doc.add_heading("Certifications", 1)
                
                for cert in cv_data["certifications"]:
                    name = cert.get("name", "")
                    issuer = cert.get("issuer", "")
                    date = cert.get("date", "")
                    
                    p = doc.add_paragraph()
                    p.add_run(f"{name}").bold = True
                    
                    p = doc.add_paragraph()
                    p.add_run(f"Issued by {issuer}, {date}")
                    
                    if "description" in cert and cert["description"]:
                        doc.add_paragraph(cert["description"])
                    
                    doc.add_paragraph()
            
            # Save the document
            doc.save(output_path)
            
            logger.info(f"DOCX document generated at {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error generating DOCX document: {str(e)}")
            raise
    
    @classmethod
    async def generate_document(cls, cv_data: Dict[str, Any], output_path: str, format: str, template_options: Optional[Dict[str, Any]] = None) -> str:
        """Generate a document in the specified format"""
        format = format.lower()
        
        if format == "pdf":
            # Determine which PDF generator to use
            pdf_generator = template_options.get("pdf_generator", "weasyprint") if template_options else "weasyprint"
            
            if pdf_generator == "reportlab":
                return await cls.generate_pdf_reportlab(cv_data, output_path, template_options)
            else:
                return await cls.generate_pdf_weasyprint(cv_data, output_path, template_options)
                
        elif format == "docx":
            return await cls.generate_docx(cv_data, output_path, template_options)
            
        else:
            raise ValueError(f"Unsupported format: {format}") 